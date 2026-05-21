from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "End_User_Documentation_Login_Registration_Module.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
BORDER = "DADCE0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_section_title(document, text):
    document.add_heading(text, level=1)


def add_subtitle(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.color.rgb = MUTED
    run.font.size = Pt(11)


def add_key_value_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.3)
    set_cell_text(table.rows[0].cells[0], "Item", True)
    set_cell_text(table.rows[0].cells[1], "Details", True)
    set_cell_shading(table.rows[0].cells[0], LIGHT_FILL)
    set_cell_shading(table.rows[0].cells[1], LIGHT_FILL)

    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, True)
        set_cell_text(cells[1], value)

    set_table_borders(table)
    document.add_paragraph()


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("End User Documentation - Login Registration Module")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED


def build_document():
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("End User Documentation")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Login and Registration Module")
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    run.bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Django Web Application | User Guide")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = MUTED

    document.add_paragraph()
    add_key_value_table(
        document,
        [
            ("Module Name", "Login and Registration Module"),
            ("Application Type", "Django web application"),
            ("Database", "SQLite using Django's built-in User model"),
            ("Primary Users", "Registered users and administrators"),
            ("Main Purpose", "Allow users to register, log in, view a protected dashboard, and log out securely."),
        ],
    )

    add_section_title(document, "1. Introduction")
    document.add_paragraph(
        "This document explains how end users and administrators can use the Login and Registration module. "
        "The module provides a simple website flow for account registration, login, dashboard access, logout, "
        "and user management through the Django admin panel."
    )

    add_section_title(document, "2. User Roles")
    add_key_value_table(
        document,
        [
            ("Visitor", "Can open the home page and choose Login or Register."),
            ("Registered User", "Can log in, access the protected dashboard, and log out."),
            ("Administrator", "Can log in to the Django admin panel and manage users."),
        ],
    )

    add_section_title(document, "3. Website Flow")
    add_subtitle(document, "The normal user journey follows this sequence:")
    for step in [
        "Open the home page.",
        "Click Register to create a new account.",
        "After successful registration, go to the Login page.",
        "Enter username and password.",
        "After successful login, view the dashboard page.",
        "Click Logout to end the session and return to the Login page.",
    ]:
        add_number(document, step)

    add_section_title(document, "4. Home Page")
    document.add_paragraph(
        "The home page is the public landing page of the application. It is accessible without logging in."
    )
    add_subtitle(document, "Available actions:")
    add_bullet(document, "Login: opens the login form for existing users.")
    add_bullet(document, "Register: opens the registration form for new users.")

    add_section_title(document, "5. User Registration")
    document.add_paragraph(
        "New users can create an account from the Register page. The form stores user details securely through Django's authentication system."
    )
    add_subtitle(document, "Registration fields:")
    add_bullet(document, "Username")
    add_bullet(document, "Email")
    add_bullet(document, "Password")
    add_subtitle(document, "Registration steps:")
    for step in [
        "Click Register on the home page or navigation bar.",
        "Enter a username, email address, and password.",
        "Submit the form.",
        "If the details are valid, the account is created and the user is redirected to the Login page.",
    ]:
        add_number(document, step)

    add_section_title(document, "6. User Login")
    document.add_paragraph(
        "Existing users can log in using their username and password. If the credentials are correct, the user is redirected to the dashboard."
    )
    add_subtitle(document, "Login steps:")
    for step in [
        "Open the Login page.",
        "Enter the registered username.",
        "Enter the password.",
        "Click Login.",
        "If the credentials are valid, the dashboard opens.",
    ]:
        add_number(document, step)
    add_subtitle(document, "Error handling:")
    add_bullet(document, "If the username or password is incorrect, an error message is displayed.")
    add_bullet(document, "The user remains on the Login page and can try again.")

    document.add_section(WD_SECTION.NEW_PAGE)
    add_section_title(document, "7. Dashboard Page")
    document.add_paragraph(
        "The dashboard is a protected page. Only logged-in users can access it. If a visitor tries to open the dashboard without logging in, the application redirects the visitor to the Login page."
    )
    add_subtitle(document, "Dashboard content:")
    add_bullet(document, "Welcome message with the logged-in username.")
    add_bullet(document, "Logout button.")
    add_bullet(document, "Admin panel button for staff users.")

    add_section_title(document, "8. Logout")
    document.add_paragraph(
        "The logout function securely ends the user's active session and redirects the user back to the Login page."
    )
    add_subtitle(document, "Logout steps:")
    for step in [
        "Click Logout from the dashboard or navigation bar.",
        "The active session is ended.",
        "The Login page opens.",
    ]:
        add_number(document, step)

    add_section_title(document, "9. Admin Panel")
    document.add_paragraph(
        "Administrators can manage users through Django's built-in admin panel. Admin access requires a superuser account."
    )
    add_key_value_table(
        document,
        [
            ("Admin URL", "/admin/"),
            ("Required Account", "Django superuser or staff account"),
            ("Main Use", "View, add, edit, and delete users"),
        ],
    )
    add_subtitle(document, "Admin tasks:")
    add_bullet(document, "View registered users.")
    add_bullet(document, "Add new users manually.")
    add_bullet(document, "Edit user details such as username, email, staff status, and permissions.")
    add_bullet(document, "Delete users when required.")

    add_section_title(document, "10. Data Storage")
    document.add_paragraph(
        "The application uses Django's built-in User model. User records are stored in the SQLite database table named auth_user."
    )
    add_subtitle(document, "Important user data:")
    add_bullet(document, "Username")
    add_bullet(document, "Email address")
    add_bullet(document, "Encrypted password")
    add_bullet(document, "Staff and admin status")

    add_section_title(document, "11. Security Notes")
    add_bullet(document, "Passwords are not stored as plain text. Django stores password hashes securely.")
    add_bullet(document, "Dashboard access is restricted to logged-in users.")
    add_bullet(document, "Logout ends the active user session.")
    add_bullet(document, "Django admin access should be limited to trusted administrators only.")

    add_section_title(document, "12. Troubleshooting")
    add_key_value_table(
        document,
        [
            ("Cannot log in", "Check that the username and password are correct."),
            ("Registration fails", "Check required fields and make sure the email is not already registered."),
            ("Dashboard redirects to login", "The user is not logged in or the session has expired."),
            ("Admin panel access denied", "The account may not have staff or superuser permissions."),
            ("Bad Request on deployed site", "Check the deployed ALLOWED_HOSTS setting."),
        ],
    )

    add_section_title(document, "13. End User Summary")
    document.add_paragraph(
        "The Login and Registration module provides a simple and secure authentication flow. Users can register, "
        "log in, access a protected dashboard, and log out. Administrators can manage user data through the Django admin panel."
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
